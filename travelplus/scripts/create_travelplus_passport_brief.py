from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "TravelPlus_Passport_Co_Che_Van_Hanh.docx"

NAVY = "073B56"
BLUE = "0087BC"
BLUE_DARK = "075D7E"
GOLD = "D99A00"
GOLD_LIGHT = "FFF5D7"
BLUE_LIGHT = "EAF7FC"
INK = "132D3B"
MUTED = "5E7280"
GRID = "D5E3EA"
GRAY_FILL = "F4F7F9"
WHITE = "FFFFFF"
GREEN = "347047"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size: float = 11, color: str = INK, bold: bool | None = None,
                 italic: bool | None = None, name: str = "Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: str = "6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_cell_text(cell, size=9.4, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        for run in p.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int],
              header_fill=BLUE_DARK, first_col_bold=False, font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    header = table.rows[0]
    repeat_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.text = text
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        style_cell_text(cell, size=9.2, color=WHITE, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, row_values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for col_idx, text in enumerate(row_values):
            cell = cells[col_idx]
            cell.text = text
            set_cell_border(cell)
            if row_idx % 2 == 1:
                set_cell_shading(cell, GRAY_FILL)
            style_cell_text(
                cell,
                size=font_size,
                color=INK,
                bold=first_col_bold and col_idx == 0,
                align=WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, len(headers) - 1) else WD_ALIGN_PARAGRAPH.CENTER,
            )
    set_table_geometry(table, widths)
    return table


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_body(doc, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_numbered(doc, title: str, detail: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    r1 = p.add_run(title + ": ")
    set_run_font(r1, bold=True, color=NAVY)
    r2 = p.add_run(detail)
    set_run_font(r2)
    return p


def add_callout(doc, label: str, text: str, fill: str = BLUE_LIGHT, accent: str = BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size="8")
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label.upper() + "\n")
    set_run_font(r1, size=9, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11.2, color=NAVY, bold=True)
    set_table_geometry(table, [9360], indent_dxa=120)
    return table


def add_spacer(doc, points=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)
    return p


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE_DARK, 15, 7),
        ("Heading 2", 13, BLUE_DARK, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("TRAVELPLUS PASSPORT  |  TÀI LIỆU NỘI BỘ")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Travel Plus  •  Cơ chế vận hành Dặm Hành Trình  •  18/08/2026")
    set_run_font(fr, size=8.2, color=MUTED)

    return doc


def build_document():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = configure_document()

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    kr = kicker.add_run("TRAVELPLUS PASSPORT")
    set_run_font(kr, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    tr = title.add_run("CƠ CHẾ VẬN HÀNH\nDẶM HÀNH TRÌNH & HẠNG THÀNH VIÊN")
    set_run_font(tr, size=23, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("Bản tóm tắt dành cho trao đổi nội bộ và phê duyệt chính sách")
    set_run_font(sr, size=11, color=MUTED, italic=True)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    mr = meta.add_run("Phiên bản 1.0  |  Ngày cập nhật: 18/08/2026  |  Đơn vị: Travel Plus")
    set_run_font(mr, size=9.3, color=MUTED, bold=True)

    add_callout(
        doc,
        "Thông điệp quản trị",
        "TravelPlus Passport không phải chương trình giảm giá đại trà. Đây là cơ chế ghi nhận khách quay lại: đặt tour để tích Dặm, dùng Dặm đổi quyền lợi và đạt hạng cao hơn để nhận ưu đãi ổn định.",
    )

    add_heading(doc, "1. Mục tiêu chương trình", 1)
    add_bullet(doc, "Tăng tỷ lệ khách quay lại đặt tour và giảm phụ thuộc vào khuyến mãi ngắn hạn.")
    add_bullet(doc, "Tạo giá trị dễ hiểu sau mỗi booking: Dặm nhận được, voucher có thể đổi và quyền lợi theo hạng.")
    add_bullet(doc, "Kiểm soát chi phí bằng tỷ lệ tích Dặm, ngưỡng voucher, trần giảm giá và điều kiện booking.")

    add_heading(doc, "2. Cơ chế vận hành tổng quát", 1)
    overview_rows = [
        ["Tên chương trình", "TravelPlus Passport - Dặm Hành Trình", "Nhận diện riêng cho chính sách thành viên"],
        ["Cách tích Dặm", "10.000đ giá trị tour đủ điều kiện = 1 Dặm", "Khuyến khích khách tiếp tục đặt tour"],
        ["Thời điểm cộng", "Sau khi booking đã thanh toán và đủ điều kiện ghi nhận", "Hạn chế booking ảo, hoàn hoặc hủy"],
        ["Dặm khả dụng", "Dùng để đổi voucher tour", "Tạo phần thưởng ngắn hạn"],
        ["Dặm xét hạng", "Dùng để xác định hạng thành viên", "Ghi nhận mức độ gắn bó dài hạn"],
        ["Khi đổi voucher", "Chỉ trừ Dặm khả dụng; không giảm Dặm xét hạng", "Khách không sợ tụt hạng khi dùng quyền lợi"],
        ["Hạn voucher", "180 ngày; dùng một lần", "Tạo động lực quay lại trong chu kỳ phù hợp"],
        ["Tại checkout", "Giảm theo hạng trước; khách chọn thêm tối đa 1 voucher", "Quyền lợi rõ ràng, ngân sách có kiểm soát"],
        ["Hoàn/hủy booking", "Thu hồi Dặm và quyền lợi phát sinh nếu không còn đủ điều kiện", "Ngăn lợi dụng chương trình"],
    ]
    add_table(doc, ["Hạng mục", "Cách vận hành", "Ý nghĩa kinh doanh"], overview_rows,
              [1900, 4000, 3460], first_col_bold=True, font_size=9.0)

    doc.add_page_break()

    add_heading(doc, "3. Hạng thành viên và quyền lợi", 1)
    add_body(doc, "Mọi thành viên được phục vụ theo cùng tiêu chuẩn Travel Plus. Hạng Passport chỉ xác định mức ưu đãi tài chính và voucher chào hạng.")
    tier_rows = [
        ["Thành viên", "0", "Chưa áp dụng", "-", "Tham gia miễn phí"],
        ["Bạc", "5.000", "1%", "200.000đ", "100.000đ\nBooking từ 3 triệu"],
        ["Vàng", "20.000", "1,5%", "400.000đ", "200.000đ\nBooking từ 6 triệu"],
        ["Kim Cương", "60.000", "2%", "600.000đ", "300.000đ\nBooking từ 10 triệu"],
        ["Signature", "150.000", "3%", "1.000.000đ", "500.000đ\nBooking từ 15 triệu"],
    ]
    add_table(doc, ["Hạng", "Dặm xét hạng", "Giảm mỗi tour", "Trần giảm", "Voucher chào hạng"],
              tier_rows, [1500, 1650, 1600, 1550, 3060], first_col_bold=True, font_size=9.0)

    add_heading(doc, "4. Bảng đổi Dặm lấy voucher", 1)
    voucher_rows = [
        ["500 Dặm", "50.000đ", "Dùng một lần", "180 ngày"],
        ["1.200 Dặm", "120.000đ", "Dùng một lần", "180 ngày"],
        ["2.500 Dặm", "250.000đ", "Dùng một lần", "180 ngày"],
    ]
    add_table(doc, ["Dặm cần đổi", "Giá trị voucher", "Cách dùng", "Thời hạn"], voucher_rows,
              [2200, 2300, 2500, 2360], first_col_bold=True, font_size=9.5)

    add_spacer(doc, 4)
    add_callout(
        doc,
        "Nguyên tắc tài chính",
        "Tỷ lệ đổi voucher hiện tương đương khoảng 1% giá trị tour đã tạo ra Dặm. Phần giảm theo hạng luôn có trần để bảo vệ biên lợi nhuận.",
        fill=GOLD_LIGHT,
        accent=GOLD,
    )

    add_heading(doc, "5. Ví dụ minh họa", 1)
    add_body(doc, "Khách đặt tour trị giá 25.000.000đ và đang ở hạng Vàng:")
    example_rows = [
        ["Giá tour", "25.000.000đ", "Giá trước quyền lợi thành viên"],
        ["Giảm hạng Vàng", "375.000đ", "1,5% và chưa vượt trần 400.000đ"],
        ["Dặm dự kiến nhận", "2.500 Dặm", "Tính theo 1 Dặm / 10.000đ"],
        ["Quyền lợi sau tour", "Đủ đổi voucher 250.000đ", "Dùng cho booking tiếp theo"],
    ]
    add_table(doc, ["Nội dung", "Kết quả", "Diễn giải"], example_rows,
              [2200, 2500, 4660], first_col_bold=True, font_size=9.3)

    doc.add_page_break()

    add_heading(doc, "6. Quy trình vận hành", 1)
    add_numbered(doc, "Đặt tour", "Hệ thống hiển thị giá, Dặm dự kiến nhận và quyền lợi có thể đạt sau tour.")
    add_numbered(doc, "Thanh toán", "Giảm theo hạng được tính trước; khách có thể chọn thêm tối đa 1 voucher Passport hợp lệ.")
    add_numbered(doc, "Xác nhận đủ điều kiện", "Booking phải hoàn tất thanh toán và không thuộc trạng thái hoàn/hủy không đủ điều kiện.")
    add_numbered(doc, "Cộng Dặm", "Hệ thống cộng đồng thời Dặm khả dụng và Dặm xét hạng theo giá trị được ghi nhận.")
    add_numbered(doc, "Đổi voucher", "Khách dùng Dặm khả dụng để đổi; Dặm xét hạng không bị giảm.")
    add_numbered(doc, "Đối soát", "Khi booking hoàn hoặc hủy, hệ thống thu hồi Dặm và quyền lợi phát sinh tương ứng.")

    add_heading(doc, "7. Quy tắc kiểm soát cần thống nhất", 1)
    add_bullet(doc, "Xác định rõ giá trị nào dùng để tính Dặm: giá tour niêm yết, giá sau giảm hay số tiền thực thu.")
    add_bullet(doc, "Quy định booking/tour không được tích Dặm hoặc không áp dụng giảm theo hạng nếu có.")
    add_bullet(doc, "Một booking chỉ dùng tối đa 1 voucher Passport; tổng mức giảm không vượt giá trị tour.")
    add_bullet(doc, "Voucher không quy đổi thành tiền mặt, không hoàn lại phần chênh lệch và chỉ dùng trong thời hạn.")
    add_bullet(doc, "Quy định thời hạn duy trì hoặc xét lại hạng cần được phê duyệt trước khi truyền thông rộng rãi.")

    add_heading(doc, "8. Chỉ số nên theo dõi", 1)
    kpi_rows = [
        ["Tỷ lệ khách quay lại", "Booking lặp lại / tổng khách thành viên", "Đo tác động giữ chân"],
        ["Tỷ lệ đổi voucher", "Voucher đã dùng / voucher đã phát hành", "Đo mức hấp dẫn quyền lợi"],
        ["Chi phí loyalty", "Tổng giảm hạng + voucher / doanh thu", "Kiểm soát ngân sách"],
        ["Doanh thu theo hạng", "Doanh thu và giá trị booking trung bình từng hạng", "Đánh giá chất lượng phân tầng"],
        ["Dặm hết hạn", "Dặm hoặc voucher hết hạn / tổng phát hành", "Theo dõi trải nghiệm và nghĩa vụ dự kiến"],
    ]
    add_table(doc, ["Chỉ số", "Cách tính gợi ý", "Mục đích"], kpi_rows,
              [2200, 4200, 2960], first_col_bold=True, font_size=9.2)

    add_spacer(doc, 8)
    add_callout(
        doc,
        "Đề xuất phê duyệt",
        "Chốt 5 nội dung trước khi ra mắt chính thức: giá trị tính Dặm, điều kiện ghi nhận, thời hạn hạng, danh sách tour loại trừ và nguyên tắc hoàn/hủy.",
        fill=BLUE_LIGHT,
        accent=BLUE,
    )

    core = doc.core_properties
    core.title = "TravelPlus Passport - Cơ chế vận hành Dặm Hành Trình"
    core.subject = "Tài liệu giới thiệu nội bộ về chương trình thành viên TravelPlus Passport"
    core.author = "Travel Plus"
    core.keywords = "TravelPlus Passport, Dặm Hành Trình, loyalty, thành viên"
    core.comments = "Phiên bản nội bộ phục vụ trao đổi và phê duyệt chính sách."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
